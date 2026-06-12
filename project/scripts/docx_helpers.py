# -*- coding: utf-8 -*-
"""Shared helpers for generating practice docx files."""

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = REPO_ROOT / "附件3：课程实践报告(参考模板).docx"

# 三人小组信息
LEADER = "曾露"
MEMBER_B = "伍灵晰"
MEMBER_C = "吴芝"
TEAM_MEMBERS = f"{LEADER}-{MEMBER_B}-{MEMBER_C}"


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def find_para(doc, keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    raise ValueError(f"paragraph not found: {keyword}")


def insert_after(ref_para, text, bold=False, indent=True, size=12):
    parent = ref_para._element.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(ref_para._element) + 1, new_p)
    para = Paragraph(new_p, ref_para._parent)
    run = para.add_run(text)
    set_run_font(run, "黑体" if bold else "宋体", size, bold)
    para.paragraph_format.line_spacing = 1.5
    if indent and not bold:
        para.paragraph_format.first_line_indent = Cm(0.74)
    return para


def add_table_after(ref_para, headers, rows):
    doc = ref_para.part.document
    parent = ref_para._element.getparent()
    idx = parent.index(ref_para._element) + 1
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            tbl.rows[r].cells[c].text = str(val)
    parent.insert(idx, tbl._element)
    blank = OxmlElement("w:p")
    parent.insert(idx + 1, blank)
    return Paragraph(blank, ref_para._parent)


class DocBuilder:
    def __init__(self, title, out_path):
        shutil.copy(TEMPLATE, out_path)
        self.doc = Document(str(out_path))
        self.out_path = out_path
        self.ref = None
        self._fill_cover(title)

    def _fill_cover(self, title):
        fields = {
            4: (title, True),
            9: ("课程名称：综合测试实践（实训04）", False),
            10: ("学    院：【请填写学院】", False),
            11: ("专业/班级：【请填写专业/班级】", False),
            12: ("任课教师：【请填写教师姓名】", False),
            13: (f"学生姓名：{LEADER}", False),
            14: ("学    号：【请填写学号】", False),
            17: ("2026 年 6 月", False),
        }
        for idx, (text, center) in fields.items():
            p = self.doc.paragraphs[idx]
            p.clear()
            run = p.add_run(text)
            set_run_font(run, "黑体" if center else "宋体", 18 if center else 12, center)
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def at(self, keyword):
        self.ref = find_para(self.doc, keyword)
        return self

    def heading(self, text):
        self.ref = insert_after(self.ref, text, bold=True, indent=False)
        return self

    def para(self, text):
        self.ref = insert_after(self.ref, text)
        return self

    def bullets(self, items):
        for item in items:
            self.ref = insert_after(self.ref, f"• {item}")
        return self

    def table(self, headers, rows):
        self.ref = add_table_after(self.ref, headers, rows)
        return self

    def placeholder(self, caption):
        self.ref = insert_after(self.ref, f"【此处插入图片：{caption}】", indent=False)
        p = self.ref
        for run in p.runs:
            set_run_font(run, "宋体", 11, color=RGBColor(128, 128, 128))
            run.italic = True
        self.ref = insert_after(self.ref, "\n\n\n\n", indent=False)
        return self

    def save(self):
        self.doc.save(self.out_path)
        print(f"Generated: {self.out_path}")
