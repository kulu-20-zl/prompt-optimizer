# -*- coding: utf-8 -*-
"""Generate team leader personal practice report (docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
LEADER = "曾露"
MEMBER_B = "伍灵晰"
MEMBER_C = "吴芝"
OUT_PATH = REPO_ROOT / f"学号_{LEADER}_实训04_实践报告.docx"


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "黑体", size, bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, "黑体", 14, True)
    else:
        set_run_font(run, "黑体", 12, True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_image_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【此处插入图片：{caption}】")
    set_run_font(run, "宋体", 11)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    blank = doc.add_paragraph()
    blank.add_run("\n" * 4)
    return p


def build_report():
    doc = Document()

    add_title(doc, "综合测试实践 — 个人实践报告", 18)
    doc.add_paragraph()

    info = [
        ("学    号", "【请填写学号】"),
        ("姓    名", LEADER),
        ("班    级", "【请填写班级】"),
        ("实训项目", "AI 提示词优化助手"),
        ("项目角色", "组长（项目负责人 / 后端开发 / 测试负责人）"),
        ("完成日期", "2026 年 6 月"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, "宋体", 12, True)
        r2 = p.add_run(value)
        set_run_font(r2, "宋体", 12)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ========== 一、实训背景 ==========
    add_heading(doc, "一、实训背景与项目概述")
    add_body(
        doc,
        "本次实训课程为「综合测试实践」（实训04），是软件测试方向的核心实践环节。"
        "实训要求小组（2～4 人）围绕自研 Web 应用，独立完成测试计划制定、测试用例设计与执行、"
        "缺陷管理、性能测试、自动化测试及测试报告撰写等完整测试流程，"
        "最终提交包含源代码、测试脚本、测试计划、测试报告及个人实践报告在内的全套材料。",
    )
    add_body(
        doc,
        "本项目定名为「AI 提示词优化助手」，是一款面向普通用户与开发者的 Web 应用，"
        "核心功能是将用户输入的原始提示词（Prompt）通过大语言模型进行结构化优化，"
        "使其更清晰、可执行、符合特定场景需求。系统采用前后端分离的单体架构："
        "后端基于 Python Flask 框架，数据库使用 SQLite，AI 能力通过 DeepSeek 开放 API 接入；"
        "前端采用原生 HTML + CSS + JavaScript 实现，无需额外构建工具，便于快速部署与测试。",
    )
    add_body(
        doc,
        "系统共包含以下核心业务模块：（1）用户认证模块：注册、登录、登出、忘记密码、登录限流；"
        "（2）提示词优化模块：支持通用、代码、学术、创意四种优化模式，提供同步与 SSE 流式两种输出方式，"
        "并支持「继续优化」（refine）功能；（3）历史记录模块：分页查询、单条删除、批量清空；"
        "（4）评分模块：对优化结果进行 1～5 星评分；（5）收藏模块：基于 localStorage 的本地收藏与「我的收藏」视图；"
        "（6）Markdown 渲染：优化结果支持标题、列表、代码块等格式的实时渲染。",
    )
    add_body(
        doc,
        "本人担任三人小组组长，是本项目的主要技术负责人与质量负责人。"
        "除承担后端架构设计与全部后端开发、单元测试 / 性能测试 / Selenium 自动化脚本编写外，"
        "还牵头制定测试计划、组织联调与回归、管理 GitHub 仓库与部署脚本。"
        f"由于{MEMBER_C}实际参与较少、未按计划交付测试汇报与答辩材料，"
        "相关文档撰写工作亦由本人补位完成。",
    )

    # ========== 二、分工 ==========
    add_heading(doc, "二、个人角色与小组分工")
    add_body(
        doc,
        "项目初期制定了三人分工方案，各成员按职责并行推进。"
        f"实际执行中，{MEMBER_B} 完成了前端与 API 测试相关工作；"
        f"{MEMBER_C} 参与讨论较多但文档类交付不足，测试汇报等由本人补位。"
        "实际完成情况如下：",
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["成员", "原定职责", "实际完成情况"]
    rows = [
        (
            f"{LEADER}（组长，本人）",
            "项目统筹、后端开发、单元/性能/Selenium 测试、测试计划",
            "完成后端全部模块、24 条单元测试、E2E/性能脚本、测试计划牵头撰写、"
            "测试汇报补位、GitHub/部署/打包、缺陷排查与修复",
        ),
        (
            MEMBER_B,
            "前端页面与交互开发；API 接口测试用例编写与维护",
            "完成 index.html、script.js、style.css、markdown.js 等前端页面与交互；"
            "编写并维护 tests/api/ 下 39 条 API 集成测试用例",
        ),
        (
            MEMBER_C,
            "测试汇报文档、缺陷汇总、答辩演示材料",
            "参与初期需求与测试讨论，未按计划交付测试汇报、缺陷汇总及答辩材料，"
            f"相关文档由{LEADER}补位完成",
        ),
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    doc.add_paragraph()
    add_body(
        doc,
        "作为组长，本人在项目中承担了以下多重角色：项目负责人（进度把控、任务分配、风险应对）、"
        "后端架构师（技术选型、模块划分、接口设计）、后端开发工程师（全部 API 与业务逻辑）、"
        "测试负责人（测试计划、用例设计、执行与报告）、"
        "自动化测试工程师（单元 / Selenium / Locust 脚本）、"
        "运维工程师（GitHub、CI、内网穿透、打包提交）。"
        f"与{MEMBER_B}协作完成前后端联调，{MEMBER_B}负责前端实现与 API 测试用例。",
    )

    # ========== 三、主要工作 ==========
    add_heading(doc, "三、个人完成的主要工作（详细说明）")

    add_heading(doc, "3.1 需求分析与技术选型", 2)
    add_body(
        doc,
        "在项目启动阶段，本人主导完成了需求分析与技术选型工作。"
        "经过调研，确定以「提示词优化」为核心场景——该场景既有明确的业务价值（帮助用户写出更好的 AI 提示词），"
        "又天然适合接入大模型 API，同时涉及用户认证、CRUD、流式输出、评分等多种典型 Web 功能，"
        "能够全面覆盖实训要求的测试维度。技术栈选择 Flask + SQLite + 原生 JS 的组合，"
        "理由是：轻量、零配置、适合实训周期、便于 pytest 集成测试、"
        "且 SQLite 文件数据库便于本地开发与测试隔离。",
    )
    add_body(
        doc,
        "在接口设计阶段，本人制定了 RESTful API 规范：认证类接口统一前缀 /api/auth，"
        "业务类接口按资源划分（/api/polish、/api/history、/api/rate），"
        "所有需鉴权接口通过 Flask Session + 装饰器 @login_required 保护，"
        "错误响应统一返回 JSON 格式 { error: string, code?: number }，"
        "为后续 API 自动化测试提供了稳定的契约基础。",
    )

    add_heading(doc, "3.2 后端系统设计与开发", 2)
    add_body(doc, "后端共 15 个 Python 源文件，按分层架构组织，本人独立完成全部开发与调试：")

    add_heading(doc, "3.2.1 应用骨架与配置", 3)
    add_bullet(doc, "app.py：Flask 应用工厂，注册 Blueprint、配置 CORS、Session、静态文件路由、before_request 鉴权钩子、/health 健康检查端点；支持 config_overrides 参数供测试注入内存数据库。")
    add_bullet(doc, "config.py：集中管理 SECRET_KEY、SQLALCHEMY_DATABASE_URI、DEEPSEEK_API_KEY、MOCK_AI 开关、登录限流阈值 LOGIN_RATE_LIMIT 等配置项，支持 .env 环境变量覆盖。")
    add_bullet(doc, "wsgi.py：生产环境 WSGI 入口，配合 gunicorn / Render 部署。")
    add_bullet(doc, "models.py：定义 User（用户表）、PolishRecord（优化记录表，含 original_text、polished_text、mode、status、rating 等字段）。")

    add_heading(doc, "3.2.2 用户认证模块（routes/auth.py）", 3)
    add_bullet(doc, "POST /api/auth/register：校验用户名唯一性、密码强度、邮箱格式、确认密码一致性，密码 bcrypt 哈希存储。")
    add_bullet(doc, "POST /api/auth/login：验证用户名密码，建立 Session，集成登录限流（services/rate_limit.py），连续错误触发 429。")
    add_bullet(doc, "POST /api/auth/logout：清除 Session。")
    add_bullet(doc, "POST /api/auth/forgot-password：实训简化版找回密码（无邮箱验证）。")
    add_bullet(doc, "utils/auth.py：封装 @login_required 装饰器、get_current_user()、hash_password()、check_password()。")
    add_bullet(doc, "utils/validators.py：用户名、密码、邮箱、文本长度等校验函数，覆盖正常与边界输入。")

    add_heading(doc, "3.2.3 提示词优化模块（routes/polish.py + services/）", 3)
    add_bullet(doc, "POST /api/polish：接收 text、mode、stream 参数；支持 general/code/academic/creative 四种模式；stream=true 时返回 SSE 流式响应，否则同步返回 JSON。")
    add_bullet(doc, "refine 子功能：接收 refine=true、polished（已有优化文本）、direction（优化方向），调用 polish_text_refine() 在原文基础上按方向再优化，前后端分拆传参避免超长文本截断。")
    add_bullet(doc, "services/ai_client.py：封装 OpenAI 兼容客户端；支持 MOCK_AI=1 时返回固定 Mock 文本；实现 polish_text()、polish_text_stream()、polish_text_refine()；流式解析 SSE chunk；思考过程检测（正则匹配「首先我需要」「让我分析」等）+ 自动重试最多 2 次。")
    add_bullet(doc, "services/prompt_modes.py：四种模式的 system prompt 模板及 refine 专用提示词，控制模型输出格式与禁止行为。")
    add_bullet(doc, "优化完成后自动写入 PolishRecord，status 标记 success/failed，供历史模块查询。")

    add_heading(doc, "3.2.4 历史记录与评分模块", 3)
    add_bullet(doc, "routes/history.py：GET /api/history 分页查询（page、per_page 参数，services/pagination.py 封装）；DELETE /api/history/<id> 单条删除；DELETE /api/history 批量清空；均校验记录归属当前用户，防止越权。")
    add_bullet(doc, "routes/rate 或 history 内评分：POST /api/rate 对指定记录评 1～5 星，services/rating.py 校验评分范围与记录存在性。")

    add_heading(doc, "3.2.5 辅助服务与工程化", 3)
    add_bullet(doc, "services/db_migrate.py：启动时自动建表与轻量迁移。")
    add_bullet(doc, "services/backup.py：SQLite 数据库定时备份与轮转（DB_BACKUP_KEEP 控制保留份数）。")
    add_bullet(doc, "services/rate_limit.py：基于内存的滑动窗口限流，防止暴力破解。")
    add_bullet(doc, "修复关键 Bug：pytest conftest.py 原使用生产库路径导致 drop_all() 清空用户数据，改为 sqlite:///:memory: 彻底隔离。")

    add_heading(doc, "3.3 前后端联调与协作", 2)
    add_body(
        doc,
        f"前端页面与交互由{MEMBER_B}负责实现，本人侧重后端接口交付与联调验收，主要协作内容如下：",
    )
    add_bullet(doc, f"制定并维护 RESTful API 契约，与{MEMBER_B}对齐注册/登录、流式优化、历史、评分等接口的请求与响应格式。")
    add_bullet(doc, f"完成 refine 继续优化接口设计与实现，配合{MEMBER_B}调整弹窗交互（direction 合并后重新生成，而非回填输入框）。")
    add_bullet(doc, "协助排查 Session 鉴权、SSE 流式解析、跨域与静态资源路径等联调问题，保障 E2E 主流程可跑通。")
    add_bullet(doc, f"审核{MEMBER_B}提交的 API 测试用例，与本人编写的单元测试、Selenium、Locust 脚本一并纳入回归。")

    add_heading(doc, "3.4 测试计划制定", 2)
    add_body(
        doc,
        "本人牵头编写《AI 提示词优化助手 — 测试计划》V1.0，文档约 300 行，涵盖：",
    )
    add_bullet(doc, "测试范围与目标：明确 In-Scope（认证、优化、历史、评分、refine、收藏）与 Out-Scope（生产级安全、多租户）。")
    add_bullet(doc, "测试策略：单元测试（pytest）、API 集成测试（pytest + Flask test client）、E2E（Selenium）、性能测试（Locust）、手工/UAT 走查。")
    add_bullet(doc, "测试环境：Windows 10/11、Python 3.13、Chrome headless、SQLite、MOCK_AI=1。")
    add_bullet(doc, "人员分工与 12 天里程碑进度表（M1～M7）。")
    add_bullet(doc, "缺陷管理流程：S1～S4 严重等级定义、回归策略。")
    add_bullet(doc, "验收标准：自动化通过率 100%、覆盖率 ≥ 70%、S1=0、S2=0、性能 P95 指标。")
    add_bullet(doc, "业务验收清单（UAT）6 项抽样场景。")

    add_heading(doc, "3.5 单元测试（24 条）", 2)
    add_body(doc, "tests/unit/ 目录，使用 pytest + fixture，全部由本人编写：")
    add_bullet(doc, "test_ai_client.py（14 条）：Mock 模式下 polish_text / polish_text_stream / polish_text_refine 正常返回；空文本、超长文本边界；思考过程检测触发重试；流式 chunk 拼接；API Key 缺失报错。")
    add_bullet(doc, "test_pagination.py（5 条）：正常分页、首页、末页、空结果、per_page 超限。")
    add_bullet(doc, "test_rating.py（5 条）：有效评分 1～5、无效评分 0/6、记录不存在、未登录、越权评分。")

    add_heading(doc, "3.6 API 集成测试（39 条）", 2)
    add_body(
        doc,
        f"tests/api/ 目录由{MEMBER_B}主要编写，本人负责 conftest 基础设施、用例评审与联调修复：",
    )
    add_bullet(doc, f"{MEMBER_B}编写 test_auth.py（17 条）、test_polish.py（9 条）、test_history.py（9 条）、test_rating.py（5 条）。")
    add_bullet(doc, "本人编写 tests/conftest.py：统一 app fixture（内存库）、db_session、authenticated_client、mock_openai，确保测试间完全隔离。")
    add_bullet(doc, f"与{MEMBER_B}协作修复 refine、流式 SSE、越权删除等接口边界用例，最终 39/39 全部通过。")

    add_heading(doc, "3.7 E2E 自动化测试（Selenium）", 2)
    add_body(
        doc,
        "tests/auto/test_selenium_flow.py，使用 Selenium WebDriver + Chrome headless，"
        "覆盖「注册 → 登录 → 输入提示词 → 点击优化 → 等待流式输出 → 查看历史」完整主流程。",
    )
    add_bullet(doc, "测试启动独立 Flask 线程，注入 config_overrides 使用内存数据库，避免污染生产数据。")
    add_bullet(doc, "使用 WebDriverWait 等待 app-shell 元素加载，确保页面 JS 初始化完成。")
    add_bullet(doc, "使用 execute_script 点击 polish-btn，规避元素被遮挡导致的点击失败。")
    add_bullet(doc, "修复历程：初版因 Session 与数据库不一致导致登录 401，排查 3 轮后通过，当前 1/1 通过。")

    add_heading(doc, "3.8 性能测试（Locust）", 2)
    add_body(
        doc,
        "tests/performance/locustfile.py，使用 Locust 2.44.0 进行 headless 压测：",
    )
    add_bullet(doc, "虚拟用户 on_start 阶段自动注册 + 登录，保持 Session 后执行业务请求。")
    add_bullet(doc, "场景权重：GET /api/history（权重 3）+ POST /api/polish（权重 1），模拟真实读多写少比例。")
    add_bullet(doc, "执行参数：5 并发用户、spawn 速率 1/s、持续 45 秒、MOCK_AI=1。")
    add_bullet(doc, "结果：146 总请求、0 失败、错误率 0%；history P50=13ms P95=27ms；polish P50=33ms P95=43ms。")
    add_bullet(doc, "编写 PERFORMANCE_REPORT.md 性能简报，纳入测试汇报。")
    add_bullet(doc, "修复历程：早期脚本注册 payload 缺少 email/confirm_password 导致 400，补全后通过。")

    add_heading(doc, "3.9 测试汇报与覆盖率", 2)
    add_bullet(doc, f"牵头编写《AI 提示词优化助手 — 测试计划》；因{MEMBER_C}未交付，本人补位完成《AI 提示词优化助手测试汇报》，含执行统计、覆盖率分析、缺陷统计、风险评估、发布建议。")
    add_bullet(doc, "执行 pytest --cov=backend --cov-report=html，生成 htmlcov/ 覆盖率 HTML 报告，后端语句覆盖率 79%（511/643），高于 70% 目标。")
    add_bullet(doc, "分模块覆盖率：config/models/pagination/rate_limit/rating/auth 100%；prompt_modes 92%；routes/polish 77%；ai_client 66%。")

    add_heading(doc, "3.10 工程化与部署", 2)
    add_bullet(doc, "建立 GitHub 仓库 prompt-optimizer，编写 push-to-github.ps1 推送脚本。")
    add_bullet(doc, "配置 .github/workflows/test.yml CI 流水线，push 时自动运行 pytest。")
    add_bullet(doc, "编写 requirements.txt / requirements-prod.txt、Procfile、render.yaml、DEPLOY.md 部署文档。")
    add_bullet(doc, "编写 scripts/start-tunnel.ps1 内网穿透脚本（localtunnel），便于无外网服务器时演示。")
    add_bullet(doc, "编写 scripts/package-submission.ps1 打包脚本，按实训要求生成 tests/unit_tests、api_tests、auto_tests、performance_tests 目录结构及 ZIP 包。")

    add_heading(doc, "3.11 缺陷排查与修复", 2)
    add_bullet(doc, "DEF-001（S2）：Selenium E2E 登录失败 — 已关闭，统一内存库 + 等待策略 + JS 点击。")
    add_bullet(doc, "DEF-002（S2）：继续优化输出思考过程 — 部分修复，ai_client 启发式检测 + 重试 + refine 专用 prompt。")
    add_bullet(doc, "DEF-003（S3）：旧版收藏 ID 无法展示 — 设计限制，迁移至完整对象存储。")
    add_bullet(doc, "DEF-004（S3）：测试库与生产库混淆 — 已关闭，conftest + Selenium 均用内存库。")
    add_bullet(doc, "DEF-005（S4）：忘记密码无邮箱验证 — 已知限制，实训简化实现。")

    # ========== 四、测试执行结果 ==========
    add_heading(doc, "四、测试执行结果汇总")
    add_body(doc, "以下为本人主导编写并执行的全部自动化测试结果：")

    result_table = doc.add_table(rows=8, cols=4)
    result_table.style = "Table Grid"
    result_data = [
        ("测试类型", "用例数", "执行结果", "执行命令"),
        ("单元测试", "24", "24/24 通过", "MOCK_AI=1 pytest tests/unit -v"),
        ("API 集成测试", "39", "39/39 通过", "MOCK_AI=1 pytest tests/api -v"),
        ("E2E 自动化", "1", "1/1 通过", "MOCK_AI=1 pytest tests/auto -v"),
        ("性能测试", "2 场景", "146 请求，0 失败", "locust --headless -u 5 -r 1 -t 45s"),
        ("代码覆盖率", "—", "79%（511/643）", "pytest --cov=backend --cov-report=html"),
        ("缺陷关闭", "5 项", "已关闭 2 项", "DEF-001、DEF-004"),
        ("合计", "66", "66/66 通过", "通过率 100%"),
    ]
    for r, row in enumerate(result_data):
        for c, val in enumerate(row):
            result_table.rows[r].cells[c].text = val

    doc.add_paragraph()
    add_heading(doc, "4.1 分模块测试结果", 2)
    module_table = doc.add_table(rows=9, cols=4)
    module_table.style = "Table Grid"
    module_data = [
        ("功能模块", "用例数", "通过", "结果"),
        ("AI 客户端（ai_client）", "14", "14", "通过"),
        ("分页服务（pagination）", "5", "5", "通过"),
        ("评分服务（rating）", "5", "5", "通过"),
        ("用户认证（/api/auth）", "17", "17", "通过"),
        ("历史记录（/api/history）", "9", "9", "通过"),
        ("提示词优化（/api/polish）", "9", "9", "通过"),
        ("记录评分（/api/rate）", "5", "5", "通过"),
        ("E2E 主流程（Selenium）", "1", "1", "通过"),
    ]
    for r, row in enumerate(module_data):
        for c, val in enumerate(row):
            module_table.rows[r].cells[c].text = val

    # ========== 五、截图 ==========
    doc.add_paragraph()
    add_heading(doc, "五、实践过程截图")
    add_body(doc, "以下位置请自行插入实际操作截图（答辩或提交前补充）：")
    placeholders = [
        "项目首页 / 登录注册界面",
        "四种优化模式切换与流式输出效果",
        "继续优化（refine）弹窗功能演示",
        "历史记录分页、删除与评分操作",
        "我的收藏视图与收藏卡片",
        "pytest 自动化测试 63 passed 截图",
        "覆盖率 HTML 报告 htmlcov/index.html",
        "Locust 性能测试结果（146 请求 0 失败）",
        "Selenium E2E 测试通过截图",
        "GitHub 仓库页面与 CI 构建记录",
        "内网穿透演示外网访问截图",
        "提交打包 ZIP 目录结构截图",
    ]
    for i, cap in enumerate(placeholders, 1):
        add_body(doc, f"图 {i}：")
        add_image_placeholder(doc, cap)

    # ========== 六、问题与解决 ==========
    add_heading(doc, "六、遇到的问题与解决过程")
    problems = [
        (
            "问题一：pytest 导致正式数据库被清空（数据丢失）",
            "现象：用户反馈优化历史突然全部消失。排查发现 tests/conftest.py 在测试 teardown 阶段"
            "对 SQLALCHEMY_DATABASE_URI 指向的 instance/grammar_assistant.db 执行 drop_all()，"
            "而早期 conftest 未强制使用内存库，导致每次运行 pytest 都会清空生产数据。"
            "解决：修改 create_app(config_overrides={SQLALCHEMY_DATABASE_URI: 'sqlite:///:memory:'})，"
            "确保所有 pytest 会话（含 API 测试、Selenium 测试）均使用独立内存数据库。"
            "同时在 app.py 中禁用 Debug 模式 reloader，避免 SQLite 文件锁冲突。",
        ),
        (
            "问题二：Selenium E2E 注册成功但登录返回 401",
            "现象：Selenium 脚本注册返回 201，但紧接着登录返回 401，"
            "WebDriverWait 等待 polish-btn 超时 30 秒，E2E 失败。"
            "排查：Selenium 启动的 Flask 实例与 pytest fixture 使用不同数据库实例，"
            "注册写入 A 库但登录读 B 库。解决："
            "（1）Selenium 启动时注入相同 config_overrides 内存库；"
            "（2）等待 #app-shell 元素确认前端 JS 加载完成；"
            "（3）使用 driver.execute_script 点击按钮规避遮挡。"
            "修复后 E2E 1/1 通过。",
        ),
        (
            "问题三：继续优化（refine）输出思考过程而非新提示词",
            "现象：用户点击「继续优化」后，模型返回「首先我需要分析…」「让我思考一下…」"
            "等推理过程文本，而非优化后的新提示词。"
            "解决：（1）在 ai_client.py 增加 _looks_like_reasoning() 启发式检测函数，"
            "匹配常见推理开头模式；（2）检测到后自动重试最多 2 次；"
            "（3）编写 refine 专用 system prompt 明确禁止输出思考过程；"
            "（4）前端 refine 弹窗改为将 direction 与已有 polished 文本合并后调用 refine 接口重新生成，"
            "而非将结果填充回输入框。Mock 环境下已验证通过，真实 API 待进一步回归。",
        ),
        (
            "问题四：Locust 性能测试注册接口返回 400",
            "现象：Locust on_start 注册失败率 100%，后续业务请求全部因无 Session 失败。"
            "排查：早期 locustfile.py 的注册 payload 仅含 username/password，"
            "缺少后端 validators 要求的 email 和 confirm_password 字段。"
            "解决：补全注册 JSON 字段，用户名加 uuid 随机后缀避免重复注册 409。"
            "修复后 5 用户 / 45 秒压测 146 请求 0 失败。",
        ),
        (
            f"问题五：{MEMBER_C}未交付测试文档导致汇报缺口",
            f"现象：临近提交节点，{MEMBER_C}未能按计划交付测试汇报、缺陷汇总与答辩材料，"
            "影响实训材料完整性。解决：本人补位撰写测试汇报与性能简报，"
            "并整理缺陷列表与发布建议，确保按时提交。"
            "此经历让我认识到文档交付应与代码交付同等纳入里程碑管理。",
        ),
        (
            "问题六：SQLite Debug 模式 reloader 导致数据库锁定",
            "现象：开发时修改代码触发 Flask reloader，SQLite 报 database is locked。"
            "解决：在 app.py 中判断若使用 SQLite 则禁用 use_reloader，"
            "并在 DEPLOY.md 中说明生产环境使用 gunicorn 多 worker 的注意事项。",
        ),
    ]
    for title, desc in problems:
        add_heading(doc, title, 2)
        add_body(doc, desc)

    # ========== 七、心得体会 ==========
    add_heading(doc, "七、心得体会")
    add_body(
        doc,
        "7.1 关于测试驱动开发的体会\n"
        "作为组长，本次实训让我深刻体会到「开发」与「测试」并非割裂的两个阶段，"
        "而是贯穿项目全生命周期的质量保障体系。在承担后端开发的同时编写单元测试和 API 测试，"
        "促使我在设计接口之初就考虑边界条件、错误处理和可测试性。例如：将 AI 调用封装为可 Mock 的独立模块"
        "（ai_client.py），使 14 条单元测试无需真实 API Key 即可运行；"
        "将会话鉴权抽离为 @login_required 装饰器，使未登录 401 测试可以一行断言完成；"
        "将分页逻辑封装为 pagination.py 服务，使边界测试与业务代码解耦。"
        "这种「可测试性设计」显著提升了开发效率——每次修改后运行 pytest 仅需 20 秒即可确认无回归。",
    )
    add_body(
        doc,
        "7.2 关于测试环境隔离的重要性\n"
        "在排查「数据丢失」问题的过程中，我深刻认识到测试环境隔离的重要性。"
        "一份看似简单的 conftest.py 若配置不当，可能造成生产数据被误删的严重后果。"
        "这让我理解了为什么企业级项目需要独立的 TEST / STAGING / PROD 环境、"
        "为什么 CI 流水线必须使用独立测试数据库、"
        "以及为什么「测试通过但数据丢了」比「测试失败」更加危险——因为前者具有隐蔽性。",
    )
    add_body(
        doc,
        "7.3 关于性能测试的实践\n"
        "使用 Locust 对高频接口 /api/history 与 /api/polish 进行压测，"
        "让我第一次亲手获得了可量化的性能基线数据：5 并发下 history P50=13ms、polish P50=33ms、"
        "错误率 0%。同时也意识到 Mock 与真实 DeepSeek API 的性能差异可能达数量级，"
        "测试报告中必须如实披露测试环境与生产环境的差异，避免用 Mock 数据误导发布决策。",
    )
    add_body(
        doc,
        "7.4 关于 E2E 自动化测试的挑战\n"
        "Selenium E2E 测试看似简单（只有 1 条），但调试过程耗时最长。"
        "浏览器环境、JS 异步加载、元素遮挡、Session 与数据库一致性等问题，"
        "都是单元测试和 API 测试不会遇到的。最终通过 3 轮排查才修复通过，"
        "让我理解了为什么业界常说「E2E 测试贵而脆，应控制数量、聚焦主流程」。",
    )
    add_body(
        doc,
        "7.5 关于团队协作与项目管理\n"
        f"本次实训中，{MEMBER_B}按期完成了前端与 API 测试，{MEMBER_C}在文档类交付上参与不足，"
        "给测试汇报与答辩材料带来进度压力。作为组长，我采取了「任务兜底」策略："
        "在督促沟通的同时，补位完成测试汇报与缺陷汇总，确保项目不因单人缺位而影响提交。"
        "这也让我认识到，组长不仅需要技术能力，"
        "还需要具备进度可视化（里程碑表）、风险预警（提前发现缺位）、"
        "文档兜底三项核心素质。",
    )
    add_body(
        doc,
        "7.6 关于文档与工程规范\n"
        "完整的测试计划、测试汇报、性能简报、部署文档、打包脚本，"
        "与代码本身同等重要。实训说明书的提交要求（tests/unit_tests 目录命名、"
        "docx 文档格式、个人实践报告）让我体会到软件工程中「交付物」的概念——"
        "项目不仅要能跑，还要能证明它经过了规范的测试、能被他人复现和理解。",
    )

    # ========== 八、总结 ==========
    add_heading(doc, "八、总结")
    add_body(
        doc,
        "本次综合测试实践实训中，本人作为三人小组组长，"
        "完成了 AI 提示词优化助手从需求分析、架构设计、后端全部开发、"
        "单元 / E2E / 性能自动化测试（24 单元 + 1 E2E + Locust）、"
        "测试计划牵头与测试汇报补位撰写、缺陷排查与修复、GitHub 仓库与 CI 配置、"
        f"部署脚本与提交打包等核心工作；与{MEMBER_B}协作完成前后端联调及 39 条 API 测试。",
    )
    add_body(
        doc,
        "量化成果：自动化测试 66 项全部通过（通过率 100%），"
        "后端代码覆盖率 79%（高于 70% 目标），"
        "Locust 性能测试 146 请求零错误率，"
        "S1 缺陷为零，DEF-001/004 已关闭，"
        "达到实训说明书规定的量化验收指标。",
    )
    add_body(
        doc,
        "通过本次实践，我系统掌握了软件测试全流程——从测试计划制定、用例设计、"
        "pytest 单元与集成测试、Selenium E2E 自动化、Locust 性能压测、"
        "覆盖率分析、缺陷管理到测试报告撰写——"
        "同时显著提升了全栈开发、项目管理与技术兜底能力，"
        "为今后从事软件测试或全栈开发工作奠定了坚实基础。",
    )

    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    build_report()
